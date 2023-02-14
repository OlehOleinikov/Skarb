"""
Формування документу MS Word
"""
import datetime
import io
import re
from typing import List
import datetime

import pandas as pd
import numpy as np

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

import matplotlib.pyplot as plt
from matplotlib.pyplot import Figure

from empty_docx import _DocEditorEmpty
from xml_converter import FileProfitXML

from defines import dict_long, dict_short, service_col_names, headersdict, dict_company_types

from functools import wraps
import time


def timeit(func):
    @wraps(func)
    def timeit_wrapper(*args, **kwargs):
        start_time = time.perf_counter()
        result = func(*args, **kwargs)
        end_time = time.perf_counter()
        total_time = end_time - start_time
        # first item in the args, ie `args[0]` is `self`
        print(f'Function {func.__name__} Took {total_time:.4f} seconds')
        return result
    return timeit_wrapper


class DocEditor(_DocEditorEmpty):
    """
    Клас формування документу зі звітом про доходи (відповідно завантажених у інстанс FileProfitXML)
    """

    def __init__(self,
                 xml_inst: FileProfitXML,
                 add_years=False,
                 add_signs=False,
                 add_tab=False,
                 sub_list_text=None,
                 sub_list_table=None):
        super().__init__()
        self.add_years = add_years
        self.add_signs = add_signs
        self.add_tab = add_tab
        self.sub_list_text = sub_list_text
        self.sub_list_table = sub_list_table

        self.xml_inst = xml_inst  # посилання на результати опрацювання XML
        self.df_xml = xml_inst.df.copy()
        self.df_xml.rename(columns=service_col_names, inplace=True)  # назви колонок до більш зручних у коді

        # Формування додаткової колонки для сортування з урахуванням кварталу:
        def create_quad_col(row):
            return int(row['year']) * 10 + int(row['quad'])

        self.df_xml['year_quad'] = self.df_xml.apply(lambda row: create_quad_col(row), axis=1)

        # Визначення переліку осіб щодо яких наявні записи у завантаженому XML:
        self.persons = [x for x in self.df_xml['person'].dropna().unique().tolist() if len(x) > 6]

    def get_available_persons(self) -> List[str]:
        return self.persons

    def write_person_to_document(self, person: str):
        DocPartPerson(self, person,
                      add_years=self.add_years, add_signs=self.add_signs, add_tab=self.add_tab,
                      sub_list_text=self.sub_list_text, sub_list_table=self.sub_list_table)


class DocPartPerson:
    """
    Клас формування частини документу, що стосується окремої особи.
    Сформована частина долучається до загального файлу.
    """
    h_pers = ['year', 'quad', 'employer_id', 'employer_name', 'income', 'tax', 'desc']

    def __init__(self,
                 editor: DocEditor,
                 person,
                 add_years=False,
                 add_signs=False,
                 add_tab=True,
                 sub_list_text=None,
                 sub_list_table=None):
        self.sub_list_text = sub_list_text
        self.sub_list_table = sub_list_table
        self.editor: DocEditor = editor
        self.document: Document = editor.document
        self.person = person
        self.df: pd.DataFrame = editor.df_xml.loc[editor.df_xml['person'] == person].copy()
        self.min_quad = self.df['year_quad'].min()
        self.max_quad = self.df['year_quad'].max()
        self.min_year = self.df['year'].min()
        self.max_year = self.df['year'].max()
        assert self.max_year >= self.min_year

        # Визначення тривалості періоду за який наявні дані (щодо опрацьованої особи):
        min_quad_val = int(str(self.min_quad)[-1])
        max_quad_val = int(str(self.max_quad)[-1])
        max_year_val = int(self.max_year)
        min_year_val = int(self.min_year)
        if self.min_year == self.max_year:
            self.dur_month = max_quad_val - min_quad_val + 1
        else:
            self.dur_month = (5 - min_quad_val) + max_quad_val + ((max_year_val - min_year_val) * 4)
        self.dur_month = self.dur_month * 3  # квартали в місяці

        # Визначення середніх значень доходів (розраховується з прибутку):
        self.profit_ave_month = round(self.df['profit'].sum() / self.dur_month, 2)
        self.profit_ave_year = round(self.profit_ave_month * 12, 2)

        # Тестове представлення тривалості у місяцях (для використання у документі):
        if self.dur_month % 12 == 0:
            self.dur_text = f'{self.dur_month // 12} р.'
        else:
            self.dur_text = f'{self.dur_month // 12} р. {self.dur_month % 12} міс.'

        # Період у кварталах
        self.quad_count = self.dur_month // 3

        # Словник відповідності: код ЄДРПОУ = назва юридичної особи
        sources_list = self.df['employer_id'].dropna().unique().tolist()
        self.sources_dict = {}
        for s in sources_list:
            name = self.df.loc[self.df['employer_id'] == s]['employer_name'].unique().tolist()[0]
            self.sources_dict.update({s: name})

        self.quad_dict = {}
        self.years_dict = {}

        # Розрахунок статистики для відображення графіків
        self._count_plot_data_by_quarts()
        self._count_plot_data_by_years()

        # Заповнення документа:
        self._add_title()
        self._add_intro()
        self._add_profit_sources()
        pivot_table_data = self._pivot_tab_data()
        self._pivot_tab_add(pivot_table_data)

        if add_years:
            self._add_profit_years()
        if add_signs:
            self._add_profit_signs()
        if add_tab:
            self._add_common_table(self.df_format(self.df, self.h_pers))
        self.document.add_page_break()

    @staticmethod
    def get_cells_grid(table):
        cells = [[]]
        col_count = table._column_count
        for tc in table._tbl.iter_tcs():
            cells[-1].append(_Cell(tc, table))
            if len(cells[-1]) == col_count:
                cells.append([])
        return cells

    @timeit
    def _count_plot_data_by_years(self):
        """Підготовка даних для гістограми - доходи по роках"""
        for pos, year in enumerate(sorted(self.df['year'].dropna().unique().tolist())):
            y_profit = round(self.df.loc[self.df['year'] == year]['profit'].sum(), 2)
            y_income = round(self.df.loc[self.df['year'] == year]['income'].sum(), 2)
            y_tax = round(self.df.loc[self.df['year'] == year]['tax'].sum(), 2)
            self.years_dict.update({pos: [None, None, str(year), y_profit, y_income, y_tax]})

    @timeit
    def _count_plot_data_by_quarts(self):
        """Підготовка даних для гістограми - доходи по кварталам"""
        cur_year = int(self.min_year)
        cur_quad = int(str(self.min_quad)[-1])
        df = self.df
        for q_order in range(self.quad_count):
            q_desc = f'{cur_year} ({cur_quad}кв.)'
            q_profit = round(df.loc[(df['year'] == cur_year) & (df['quad'] == cur_quad)]['profit'].sum(), 2)
            q_income = round(df.loc[(df['year'] == cur_year) & (df['quad'] == cur_quad)]['income'].sum(), 2)
            q_tax = round(df.loc[(df['year'] == cur_year) & (df['quad'] == cur_quad)]['tax'].sum(), 2)
            self.quad_dict.update({q_order: [cur_year, cur_quad, q_desc, q_profit, q_income, q_tax]})
            cur_quad += 1
            if cur_quad == 5:
                cur_year += 1
                cur_quad = 1

    @staticmethod
    def f2s(amount: float):
        """Перетворення float у рядок string формату 1 200 000.00 (для відображення у документах)"""
        try:
            thou_sep = ' '
            deci_sep = '.'
            w_dec = '%.2f' % amount
            part_int = w_dec.split('.')[0]
            part_int = re.sub(r"\B(?=(?:\d{3})+$)", thou_sep, part_int)
            part_dec = w_dec.split('.')[1]
            return part_int + deci_sep + part_dec
        except Exception:
            print(f'Error with float value {amount} (type {type(amount)}) - cant convert to string')
            return 'n/a'

    @timeit
    def _add_title(self):
        """Друк заголовку документа"""
        self.document.add_paragraph(f'_______ (РНОКПП {str(self.person)})', style='central_header')
        self.document.add_paragraph()

    @staticmethod
    def df_format(df, headers):
        """Форматування датафрейму для відображення у документі"""
        df = df[headers].copy()

        def f2s_wrap(val):
            return DocPartPerson.f2s(val)

        df['income'] = df['income'].apply(lambda x: f2s_wrap(x))
        df['tax'] = df['tax'].apply(lambda x: f2s_wrap(x))
        # df['profit'] = df['profit'].apply(lambda x: f2s_wrap(x))

        # Зменшення кількості колонок:
        df['year'] = df['year'].astype(str) + ' (' + df['quad'].astype(str) + 'кв.)'
        df['employer_name'] = df['employer_name'].astype(str) + ' (код ' + df['employer_id'].astype(str) + ')'
        df = df[['year', 'employer_name', 'income', 'tax', 'desc']]
        df.replace({'desc': dict_long}, inplace=True)
        df.rename(columns=headersdict, inplace=True)
        df.fillna('Не зазначено', inplace=True)
        return df

    @timeit
    def _add_intro(self):
        """Друк вступний текст з загальною сумою доходу та середніми значеннями"""
        p_points_intro = self.document.add_paragraph(style='text_base')
        p_points_intro.add_run(
            f"Опрацюванням відомостей витягу Державного реєстру фізичних осіб - платників податків про суми доходів "
            f"та нарахованих податків (платник ______, РНОКПП {self.person}) за період {str(self.min_quad)[-1]}кв. "
            f"{self.min_year} року - {str(self.max_quad)[-1]}кв. {self.max_year} року (загальний період "
            f"{self.dur_text}) встановлено отримання доходів на суму {self.f2s(self.df['income'].sum())} грн., "
            f"утримано податків на суму {self.f2s(self.df['tax'].sum())} грн.")
        p_points_intro.add_run(f" (прибуток складає {self.f2s(self.df['profit'].sum())} грн.):").bold = True

        p_average_y = self.document.add_paragraph(style='List Bullet 2')
        p_average_y.add_run(f"в середньому на рік - ")
        p_average_y.add_run(f"{self.f2s(self.profit_ave_year)} грн.").bold = True

        p_average_m = self.document.add_paragraph(style='List Bullet 2')
        p_average_m.add_run(f'в середньому на місяць - ')
        p_average_m.add_run(f'{self.f2s(self.profit_ave_month)} грн.').bold = True
        p_dummy = self.document.add_paragraph('', style='text_base')

    @timeit
    def _add_plot(self, input_data: dict):
        """Графік загального прибутку по роках / кварталах"""
        columns = []  # квартал, рік...
        data = []  # двомірний масив ([[рік, квартал, заголовок, сума, сума сума],[рядок...],...)
        for i, v in input_data.items():
            v: list
            data.append([v[3], v[5], v[4]])
            columns.append(v[2])

        data_np = np.array(data).transpose()
        rows = ["Дохід", "Податок", 'Прибуток']
        n_rows = len(data_np)

        values = np.linspace(0, int(np.amax(data_np)), 5)  # положення підписів осі y
        values_lbl = np.linspace(0, int(np.amax(data_np)) // 1000, 5)  # підписи осі y

        colors = plt.cm.BuPu(np.linspace(0, 0.5, len(rows)))
        index = np.linspace(0.5, len(columns) - 0.5, len(columns))  # положення барів по осі х
        bar_width = 0.5

        cell_text = []  # значення для заповнення таблички під графіком
        for row in range(n_rows):
            cell_text.append(['%d' % (x / 1000.0) for x in data_np[row]])
        cell_text.reverse()

        fig, ax = plt.subplots()
        fig: Figure
        fig.set_dpi(100)
        fig.set_size_inches(10, 3.5, forward=True)
        ax.set_xlim(0, len(columns))

        plt.style.use('seaborn-whitegrid')
        plt.bar(index, data_np[2], bar_width, color=colors[2], edgecolor='black')
        plt.bar(index, data_np[0], bar_width, color=colors[1], edgecolor='black')

        the_table = plt.table(cellText=cell_text,
                              rowLabels=rows,
                              colLabels=columns,
                              loc='bottom')
        the_table.scale(1, 2)

        plt.ylabel("Тисяч грн.")
        plt.yticks(values, ['%d' % val for val in values_lbl])
        plt.xticks([])
        plt.subplots_adjust(bottom=0.3)
        ax.margins(x=0.0, y=0.05)
        # plt.title('Дохід поквартально')

        memory_file = io.BytesIO()
        plt.savefig(memory_file)
        p_plot_timeline = self.document.add_paragraph(style='central_header')
        p_plot_timeline.add_run().add_picture(memory_file, width=Cm(17))
        self.document.add_paragraph(style='text_base')

    @timeit
    def _add_pie(self, data_ser: pd.Series, percent_limit=5, hide_labels=False):
        # Групування малозначних записів у рядок "Інші":
        all_amount = data_ser.sum()
        limit = (all_amount / 100) * percent_limit
        rate_show = data_ser.loc[data_ser >= limit]
        rate_hide = data_ser.loc[data_ser < limit]
        hide_sum = rate_hide.sum()
        rate = pd.concat([rate_show, pd.Series(index=['Інші'], data=[hide_sum])])

        # Основні масиви для побудови кругового графіку:
        order = [str(f'№{int(x)}') for x in list(np.linspace(1, len(rate), len(rate)))]
        desc = list(rate.index)
        vals = rate.to_list()
        vals_lbl = [re.sub(r"\B(?=(?:\d{3})+$)", ' ', str(int(x/1000))) for x in vals]  # у вигляді тис. з розділювачем

        plt.style.use('seaborn-whitegrid')
        fig, ax = plt.subplots()
        fig: Figure
        fig.set_dpi(100)
        fig.set_size_inches(14, 6)
        ax.margins(x=0.0, y=0.05)

        pie_labels = order if hide_labels else desc
        patches, texts, autotexts = ax.pie(vals,
                                           labels=pie_labels,
                                           autopct='%1.0f%%',
                                           shadow=True,
                                           # colors=colors,
                                           startangle=90,
                                           frame=False,
                                           radius=1.1,
                                           wedgeprops={"edgecolor": "k", 'linewidth': 0.8},
                                           labeldistance=1.2,
                                           explode=[0.05] * len(vals))
        for t in texts:
            t.set_fontsize(24)
        for t in autotexts:
            t.set_fontsize(24)

        memory_file = io.BytesIO()
        plt.savefig(memory_file)
        p_plot_pie = self.document.add_paragraph(style='central_header')
        p_plot_pie.add_run().add_picture(memory_file, width=Cm(14))
        self.document.add_paragraph(style='text_base')

        # Додати графічну таблицю з легендою графіку:
        fig, ax = plt.subplots()
        fig: Figure
        fig.set_dpi(100)
        fig.set_size_inches(8, 0.3*len(vals_lbl))
        plt.xticks([])
        plt.yticks([])
        ax.set_frame_on(False)
        col_labels = ["Вид доходу", "Тис.грн."]
        cell_text = []
        for row in range(len(rate)):
            cell_text.append([desc[row], vals_lbl[row]])

        table = plt.table(cellText=cell_text,
                          rowLabels=order,
                          colLabels=col_labels,
                          loc='center',
                          colWidths=[0.9, 0.3],
                          cellLoc='center')
        table.auto_set_font_size(False)
        table.set_fontsize(14)
        table.scale(1, 1.3)
        # plt.subplots_adjust(bottom=0.3)
        ax.margins(x=0.0, y=0.05)

        memory_file = io.BytesIO()
        plt.savefig(memory_file)
        p_plot_pie_table = self.document.add_paragraph(style='central_header')
        p_plot_pie_table.add_run().add_picture(memory_file, width=Cm(12))
        self.document.add_paragraph(style='text_base')

    @timeit
    def _add_profit_sources(self):
        """Деталізація по джерелам доходів (місцям роботи)"""
        self.document: Document()
        p_sources = self.document.add_paragraph('', style='text_base')
        p_sources.add_run('Джерела доходів:').bold = True

        employer_rating = self.df.groupby('employer_id')['income'].sum()
        employer_rating = employer_rating.sort_values(ascending=False)
        emp_df = self._prep_emp_df(employer_rating)

        self._add_employer_table(emp_df)
        self.document.add_paragraph('', style='text_base')

    @timeit
    def _add_profit_signs(self):
        """Деталізація по видам доходів"""
        self.document: Document()
        p_signs = self.document.add_paragraph('', style='text_base')
        p_signs.add_run('Ознаки (види) доходів:').bold = True

        signs_rating = self.df.groupby('desc')['income'].sum()
        signs_rating = signs_rating.sort_values(ascending=False)

        df_short = self.df[['desc', 'income']].copy()
        df_short.replace({'desc': dict_short}, inplace=True)
        signs_rating_pie = df_short.groupby('desc')['income'].sum()
        if len(signs_rating_pie) > 1:
            self._add_pie(signs_rating_pie)

        for sign in list(signs_rating.index):
            s_p = self.document.add_paragraph(f"{self.f2s(signs_rating[sign])} грн. - {dict_long.get(sign, sign)}",
                                              style='List Bullet')
            if self.sub_list_text:
                df_sign = self.df.loc[self.df['desc'] == sign]
                employers_in_sign = df_sign.groupby('employer_id')['income'].sum()
                employers_in_sign = employers_in_sign.sort_values(ascending=False)
                if len(employers_in_sign) > 0:
                    s_p.add_run(':')
                    for cur_emp in list(employers_in_sign.index):
                        self.document.add_paragraph(f"{self.f2s(employers_in_sign[cur_emp])} грн. - код {cur_emp} "
                                                    f"({self.sources_dict.get(cur_emp, 'назва не зазначається')})",
                                                    style='List Bullet 2')
            if self.sub_list_table:
                df_sign = self.df.loc[self.df['desc'] == sign]
                employers_in_sign = df_sign.groupby('employer_id')['income'].sum()
                employers_in_sign = employers_in_sign.sort_values(ascending=False)
                if len(employers_in_sign) > 0:
                    s_p.add_run(':')
                    emp_df = self._prep_emp_df(employers_in_sign)
                    self._add_employer_table(emp_df)
        self.document.add_paragraph('', style='text_base')

    @timeit
    def _add_profit_years(self):
        """Деталізація по роках"""
        self.document: Document()
        p_years = self.document.add_paragraph('', style='text_base')
        p_years.add_run('Доходи по роках:').bold = True

        if self.dur_month > 36:  # Графік з річною деталізацією, якщо даних багато
            self._add_plot(self.years_dict)
        else:  # Графік з поквартальною деталізацією, якщо даних небагато
            self._add_plot(self.quad_dict)

        years_rating = self.df.groupby('year')['income'].sum()
        years_rating = years_rating.sort_index(ascending=False)
        for year in list(years_rating.index):
            y_p = self.document.add_paragraph(f"{year} рік - {self.f2s(years_rating[year])} грн.", style='List Bullet')

            if self.sub_list_text:
                df_year = self.df.loc[self.df['year'] == year]
                year_emps = df_year.groupby('employer_id')['income'].sum()
                if len(years_rating) > 0:
                    y_p.add_run(':')
                    year_emps = year_emps.sort_values(ascending=False)
                    for emp in list(year_emps.index):
                        self.document.add_paragraph(f"{self.f2s(year_emps[emp])} грн. - код {emp} "
                                                    f"({self.sources_dict.get(emp, 'назва не зазначається')})",
                                                    style='List Bullet 2')
            if self.sub_list_table:
                df_year = self.df.loc[self.df['year'] == year]
                year_emps = df_year.groupby('employer_id')['income'].sum()
                year_emps = year_emps.sort_values(ascending=False)
                if len(years_rating) > 0:
                    y_p.add_run(':')
                    emp_df = self._prep_emp_df(year_emps)
                    self._add_employer_table(emp_df)
        self.document.add_paragraph('', style='text_base')

    @timeit
    def _add_common_table(self, df: pd.DataFrame):
        print(f'Test common table (from pandas dataframe): {df.shape[0]} rows, {df.shape[1]} columns')
        """Додавання до документа таблиці з відомостями про всі доходи деталізовано"""
        assert len(df.columns) == 5, 'Очікується, що в загальній таблиці має бути 5 колонок'
        df = df.copy(deep=True)
        df: pd.DataFrame
        df.reset_index(inplace=True, drop=True)

        p_table_intro = self.document.add_paragraph(style='text_base')
        p_table_intro.add_run("Деталізована таблиця відомостей про отримані доходи: ")

        # Створення таблиці та заповнення кольором заголовків:
        time_point = time.perf_counter()
        tab = self.document.add_table(rows=df.shape[0] + 1, cols=len(df.columns))
        tab.allow_autofit = False
        tab.style = 'Table Grid'
        headers = list(df.columns)
        cells = self.get_cells_grid(tab)
        for pos, header in enumerate(headers):
            cells[0][pos].text = str(header)
            cell_xml_element = cells[0][pos]._tc
            table_cell_properties = cell_xml_element.get_or_add_tcPr()
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'), 'd9d9d9')
            table_cell_properties.append(shade_obj)

        print(f'\tCreate table, fill headers: {(time.perf_counter() - time_point):.4f}')

        # Внесення даних у таблицю
        time_point = time.perf_counter()
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                cells[i+1][j].text = df.values[i, j]
        print(f'\tFill data: {(time.perf_counter() - time_point):.4f}')

        # Центрування колонок
        time_point = time.perf_counter()
        for i in range(df.shape[0]):
            row_in_tab = i + 1
            cells[row_in_tab][0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cells[row_in_tab][1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            cells[row_in_tab][2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            cells[row_in_tab][3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            cells[row_in_tab][4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        print(f'\tAlign column content: {(time.perf_counter() - time_point):.4f}')

        # Формат заголовків таблиці
        for cell_pos in range(len(headers)):
            tab.rows[0].cells[cell_pos].paragraphs[0].runs[0].font.bold = True
            tab.rows[0].cells[cell_pos].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Встановлення ширини колонок
        time_point = time.perf_counter()
        widths = (Cm(2), Cm(5.5), Cm(2), Cm(2), Cm(5.5))
        for i in range(df.shape[0]+1):
            for idx, width in enumerate(widths):
                cells[i][idx].width = width
        print(f'\tSet columns width: {(time.perf_counter() - time_point):.4f}')
        self.document.add_paragraph(style='text_base')

    @timeit
    def _add_employer_table(self, df: pd.DataFrame):
        """Додавання до документу таблиці зі статистикою отриманих сум від працедавців"""
        assert len(df.columns) == 3, 'Очікується, що в таблиці працедавців має бути 3 колонки'
        df = df.copy(deep=True)
        df: pd.DataFrame
        df.reset_index(inplace=True, drop=True)

        # Створення таблиці та заповнення кольором заголовків:
        tab = self.document.add_table(rows=df.shape[0] + 1, cols=len(df.columns))
        tab.allow_autofit = False
        tab.style = 'Table Grid'
        headers = list(df.columns)
        for pos, header in enumerate(headers):
            tab.rows[0].cells[pos].text = str(header)
            cell_xml_element = tab.rows[0].cells[pos]._tc
            table_cell_properties = cell_xml_element.get_or_add_tcPr()
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'), 'd9d9d9')
            table_cell_properties.append(shade_obj)
        cells = self.get_cells_grid(tab)

        # Внесення даних у таблицю
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                cells[i+1][j].text = str(df.values[i, j])

        # Центрування колонок
        for row in range(len(tab.rows)):
            cells[row][0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            cells[row][1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cells[row][2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Формат заголовків таблиці
        for cell_pos in range(len(headers)):
            cells[0][cell_pos].paragraphs[0].runs[0].font.bold = True
            cells[0][cell_pos].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Встановлення ширини колонок
        widths = (Cm(2.5), Cm(2.5), Cm(12.0))
        for i in range(df.shape[0]+1):
            for idx, width in enumerate(widths):
                cells[i][idx].width = width
        self.document.add_paragraph(style='text_base')

    @timeit
    def _prep_emp_df(self, employer_rating: pd.Series):
        """Перетворення статистики сум по агентами у датафрейм (для подальшої побудови таблиці документу)"""
        emp_code_list = list(employer_rating.index)
        emp_vals_list = employer_rating.tolist()
        data = [emp_vals_list, emp_code_list, emp_code_list]
        emp_df = pd.DataFrame(data).transpose()
        emp_df.columns = ['Сума грн.', "Код агента", "Найменування"]
        emp_df.replace({'Найменування': self.sources_dict}, inplace=True)
        emp_df['Найменування'] = emp_df['Найменування'].apply(lambda cell: self.company_title(cell))

        def f2s_wrap(val):
            return DocPartPerson.f2s(val)
        emp_df['Сума грн.'] = emp_df['Сума грн.'].apply(lambda x: f2s_wrap(x))
        return emp_df

    @timeit
    def _pivot_tab_data(self):
        """Підготовка списку з даними для зведеної таблиці (клітинки, що мають злитись вертикально - порожні)"""
        df = self.df.copy()
        piv = pd.pivot_table(df,
                             index=['year', 'desc', 'employer_id'],
                             values=['profit'],
                             aggfunc=np.sum)
        indexes = list(piv.index)
        cells = []
        last_y = None
        last_y_s = None
        for turn in range(len(indexes)):
            row = []
            cur_y = str(indexes[turn][0])
            cur_y_s = cur_y + str(indexes[turn][1])
            if cur_y != last_y:
                row.append(cur_y)
                row.append(self.f2s(piv.loc[int(cur_y), :, :].sum()))
                # year_profit = re.sub(r"\B(?=(?:\d{3})+$)", ' ', str(int(piv.loc[int(cur_y), :, :].sum())//1000))
                last_y = cur_y
            else:
                row.append('')
                row.append('')
            if last_y_s != cur_y_s:
                row.append(f'{dict_short.get(indexes[turn][1], "Вид відсутній у довідниках")} (код {indexes[turn][1]})'
                           f' -   {"%.2f" % float(piv.loc[int(cur_y), indexes[turn][1], :].sum())} грн.')
                last_y_s = cur_y_s
            else:
                row.append('')

            row.append(f'КОД {str(indexes[turn][2])} - '
                       f'{self.company_title(self.sources_dict.get(indexes[turn][2], "(!)"))}')
                       # f'({self.f2s(piv.loc[int(cur_y), indexes[turn][1], indexes[turn][2]].sum())} грн.)'
            row[1], row[3] = row[3], row[1]
            row[1], row[2] = row[2], row[1]
            cells.append(row)
        return cells

    @staticmethod
    def company_title(full_name):
        """Застосування скорочень до найменування організаційно-правової форми юридичної особи"""
        full_name = re.sub(' +', ' ', full_name)
        for key, value in dict_company_types.items():
            full_name = re.sub(key, value.upper(), full_name, flags=re.IGNORECASE)
        return full_name

    @timeit
    def _pivot_tab_add(self, data: List[List[str]]):
        """Додавання до документу форматованої зведеної таблиці РІК - ВИД - ЮРИДИЧНА ОСОБА - СУМА ЗА РІК """
        print(f'Test pivot table: {len(data)} rows')
        time_point = time.perf_counter()
        headers = ['Рік', "Вид доходу", "Найменування агента", "Сума (грн.)"]
        p_table_intro = self.document.add_paragraph(style='text_base')
        p_table_intro.add_run("Зведена таблиця доходів в розрізі періодів та видів: ")
        print(f'\tParagraph add: {(time.perf_counter() - time_point):.4f}')

        # Створення таблиці та заповнення кольором заголовків:
        time_point = time.perf_counter()
        tab = self.document.add_table(rows=len(data) + 1, cols=len(headers))
        tab.allow_autofit = False
        tab.style = 'Table Grid'
        cells = self.get_cells_grid(tab)
        for pos, header in enumerate(headers):
            cells[0][pos].text = str(header)
            cell_xml_element = cells[0][pos]._tc
            table_cell_properties = cell_xml_element.get_or_add_tcPr()
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'), 'd9d9d9')
            table_cell_properties.append(shade_obj)

        print(f'\tСтворення таблиці та заповнення кольором заголовків: {(time.perf_counter() - time_point):.4f}')

        # Внесення даних у таблицю
        time_point = time.perf_counter()
        for i in range(len(data)):
            for j in range(len(headers)):
                cells[i+1][j].text = data[i][j]
        print(f'\tВнесення даних у таблицю: {(time.perf_counter() - time_point):.4f}')

        # Злиття клітинок:
        time_point = time.perf_counter()
        for column in range(len(headers)):
            last_filled = 0
            for row in range(len(tab.rows)):
                if cells[row][column].paragraphs[0].runs[0].text != "":
                    if (row - last_filled) > 1:
                        a = cells[last_filled][column]
                        b = cells[row-1][column]
                        a.merge(b)
                    last_filled = row

            if cells[len(data)][column].paragraphs[0].runs[0].text == "":
                a = cells[last_filled][column]
                b = cells[len(data)][column]
                a.merge(b)
        print(f'\tЗлиття клітинок: {(time.perf_counter() - time_point):.4f}')

        # Центрування колонок
        time_point = time.perf_counter()
        for row in range(len(tab.rows)):
            cells[row][0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cells[row][1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            cells[row][2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            cells[row][3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        print(f'\tЦентрування колонок: {(time.perf_counter() - time_point):.4f}')

        # Центрування клітинок по вертикалі
        time_point = time.perf_counter()
        for row in range(len(tab.rows)):
            for vertical_col in [0, 1, 3]:
                cells[row][vertical_col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells[row][vertical_col].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        print(f'\tЦентрування клітинок по вертикалі: {(time.perf_counter() - time_point):.4f}')

        # Формат років (значень першої колонки)
        time_point = time.perf_counter()
        for row in range(len(tab.rows)):
            # cells[row][0].paragraphs[0].runs[0].font.bold = True
            cells[row][0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        print(f'\tФормат років (значень першої колонки): {(time.perf_counter() - time_point):.4f}')

        # Формат заголовків таблиці
        time_point = time.perf_counter()
        for cell_pos in range(len(headers)):
            cells[0][cell_pos].paragraphs[0].runs[0].font.bold = True
            cells[0][cell_pos].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cells[0][cell_pos].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        print(f'\tФормат заголовків таблиці: {(time.perf_counter() - time_point):.4f}')

        # Встановлення ширини колонок
        time_point = time.perf_counter()
        widths = (Cm(1), Cm(4.7), Cm(8.8), Cm(2.5))
        for i in range(len(data) + 1):
            for idx, width in enumerate(widths):
                cells[i][idx].width = width
        print(f'\tSet columns\' width: {(time.perf_counter() - time_point):.4f}')

        # Видалення порожніх рядків після злиття:
        time_point = time.perf_counter()
        for column in range(len(headers)):
            for row in range(len(tab.rows)):
                cur_paragraphs = cells[row][column].paragraphs
                if len(cur_paragraphs) > 1:
                    for paragraph in cur_paragraphs[1:]:
                        if len(paragraph.text) == 0:
                            p = paragraph._element
                            p.getparent().remove(p)
                            p._p = p._element = None
        print(f'\tВидалення порожніх рядків після злиття: {(time.perf_counter() - time_point):.4f}')
        self.document.add_paragraph(style='text_base')
























