from pathlib import Path
import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class _DocEditorEmpty:
    """
    Клас, що містить інстанс порожнього документу (python-docx) з типовими налаштуваннями та методами
    """

    def __init__(self):
        self.document = Document()
        self.sections = self.document.sections
        """Розмір полів"""
        for section in self.sections:
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
            section.top_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(1)
            section.bottom_margin = Cm(2)

        """Стилі тексту"""
        self.styles = self.document.styles
        # Стиль - заголовок "по центру":
        header_c = self.styles.add_style('central_header', WD_STYLE_TYPE.PARAGRAPH)
        header_c.font.name = 'Times New Roman'
        header_c.font.size = Pt(14)
        header_c.font.bold = True
        header_c.paragraph_format.space_before = Pt(0)
        header_c.paragraph_format.space_after = Pt(0)
        header_c.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Стиль - звичайний текст:
        text_base = self.styles.add_style('text_base', WD_STYLE_TYPE.PARAGRAPH)
        text_base.font.name = 'Times New Roman'
        text_base.font.size = Pt(14)
        text_base.paragraph_format.space_before = Pt(0)
        text_base.paragraph_format.space_after = Pt(0)
        text_base.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        text_base.paragraph_format.first_line_indent = Cm(1.25)
        text_base.paragraph_format.line_spacing = 1

        # Стиль - темно-червоний текст:
        text_red = self.styles.add_style('text_red', WD_STYLE_TYPE.PARAGRAPH)
        text_red.font.name = 'Times New Roman'
        text_red.font.size = Pt(14)
        text_red.font.color.rgb = RGBColor(127, 12, 7)
        text_red.paragraph_format.space_before = Pt(0)
        text_red.paragraph_format.space_after = Pt(0)
        text_red.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        text_red.paragraph_format.first_line_indent = Cm(1.25)
        text_red.paragraph_format.line_spacing = 1

        # Дефолтний стиль - ненумерований список:
        list_style = self.styles['List Bullet']
        list_style.font.name = 'Times New Roman'
        list_style.font.size = Pt(14)
        list_style.paragraph_format.space_before = Pt(0)
        list_style.paragraph_format.space_after = Pt(0)
        list_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        list_style.paragraph_format.first_line_indent = Cm(-0.5)
        list_style.paragraph_format.left_indent = Cm(0.63)
        list_style.paragraph_format.line_spacing = 1

        # Дефолтний стиль - ненумерований список другого порядку:
        list_style = self.styles['List Bullet 2']
        list_style.font.name = 'Times New Roman'
        list_style.font.size = Pt(14)
        list_style.paragraph_format.space_before = Pt(0)
        list_style.paragraph_format.space_after = Pt(0)
        list_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        list_style.paragraph_format.first_line_indent = Cm(-0.5)
        list_style.paragraph_format.left_indent = Cm(2.5)
        list_style.paragraph_format.line_spacing = 1

        tab_style = self.styles['Table Grid']
        tab_style.font.name = 'Times New Roman'
        tab_style.font.size = Pt(9)

    # def add_tab(self, df: pd.DataFrame):
    #     """Підготовка та додавання таблиці у документ"""
    #
    #     df = df.copy(deep=True)
    #     df: pd.DataFrame
    #     df.reset_index(inplace=True, drop=True)
    #
    #     p_table_intro = self.document.add_paragraph(style='text_base')
    #     p_table_intro.add_run("Деталізована таблиця відомостей про отримані доходи: ")
    #
    #     # Перенесення даних з датафрейму у таблицю документа:
    #     tab = self.document.add_table(rows=df.shape[0] + 1, cols=len(df.columns))
    #     tab.allow_autofit = False
    #     tab.style = 'Table Grid'
    #     headers = list(df.columns)
    #     for pos, header in enumerate(headers):
    #         tab.rows[0].cells[pos].text = str(header)
    #         cell_xml_element = tab.rows[0].cells[pos]._tc
    #         table_cell_properties = cell_xml_element.get_or_add_tcPr()
    #         shade_obj = OxmlElement('w:shd')
    #         shade_obj.set(qn('w:fill'), 'd9d9d9')
    #         table_cell_properties.append(shade_obj)
    #
    #     for row_index, row in df.iterrows():
    #         pos = row_index + 1
    #         for df_cell in range(len(df.columns)):
    #             tab.rows[pos].cells[df_cell].text = str(df.iat[row_index, df_cell])
    #
    #     for cell_pos in range(len(headers)):
    #         tab.rows[0].cells[cell_pos].paragraphs[0].runs[0].font.bold = True
    #         tab.rows[0].cells[cell_pos].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #
    #     self.document.add_paragraph(style='text_base')

    def save_docx(self, file_path):
        """Збереження документу python-docx у файл MS Word"""
        if type(file_path) == str:
            file_path = Path(file_path)
            file_path = file_path.with_suffix('.docx')
        try:
            self.document.save(file_path.absolute())
            return True
        except Exception:
            return False